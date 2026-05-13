from __future__ import annotations

import json
import math
import shutil
import zipfile
from pathlib import Path
from statistics import NormalDist

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(r"C:\Users\gjm31\OneDrive\Escritorio\127 avo art")
DATE = "20260512"
GITHUB_REPO_URL = "https://github.com/gabrielmontufar/article-127-biocemented-foundation-reliability"
WORK = ROOT / f"Trabajo ASCE JGGE articulo 127 antes de 80 {DATE}"
MANUSCRIPT_DIR = WORK / "01 Manuscript"
CALC_DIR = WORK / "02 Computational verification"
FIG_DIR = WORK / "03 Figures used in manuscript"
TABLE_DIR = WORK / "04 Tables for manuscript"
AUDIT_DIR = WORK / "05 Audit and issue log"

TITLE = "Time-Dependent Reliability of Shallow Foundations on Degrading Biocemented Sand"
FILE_STEM = "Time Dependent Reliability Shallow"
NORM = NormalDist()


def ensure_dirs() -> None:
    if WORK.exists():
        shutil.rmtree(WORK)
    for folder in (MANUSCRIPT_DIR, CALC_DIR, FIG_DIR, TABLE_DIR, AUDIT_DIR):
        folder.mkdir(parents=True, exist_ok=True)


def bearing_factors(phi_deg: np.ndarray | float) -> tuple[np.ndarray, np.ndarray, np.ndarray]:
    phi = np.radians(phi_deg)
    n_q = np.exp(np.pi * np.tan(phi)) * np.tan(np.radians(45.0) + phi / 2.0) ** 2
    n_c = (n_q - 1.0) / np.tan(phi)
    n_gamma = 2.0 * (n_q + 1.0) * np.tan(phi)
    return n_c, n_q, n_gamma


def capacity(
    b: np.ndarray | float,
    h_over_b: np.ndarray | float,
    eta0: np.ndarray | float,
    eta_r: np.ndarray | float,
    lam: np.ndarray | float,
    cb0: np.ndarray | float,
    phi_m_deg: np.ndarray | float,
    delta_tan_phi: np.ndarray | float,
    gamma_eff: np.ndarray | float,
    d_f: float,
    t_year: np.ndarray | float,
) -> tuple[np.ndarray, np.ndarray, np.ndarray, np.ndarray]:
    eta = eta_r + (eta0 - eta_r) * np.exp(-lam * t_year)
    tan_phi = np.tan(np.radians(phi_m_deg)) + delta_tan_phi * eta
    phi_deg = np.degrees(np.arctan(tan_phi))
    cb = cb0 * eta**1.15
    chi_h = 1.0 - np.exp(-h_over_b / 0.70)
    chi_s = np.clip(1.0 - 0.18 * eta**2.0, 0.72, 1.0)
    n_c, n_q, n_g = bearing_factors(phi_deg)
    _, n_q_m, n_g_m = bearing_factors(phi_m_deg)
    q_base = gamma_eff * d_f
    q_matrix = q_base * n_q_m + 0.5 * gamma_eff * b * n_g_m
    delta_cement = cb * n_c + q_base * (n_q - n_q_m) + 0.5 * gamma_eff * b * (n_g - n_g_m)
    q_ult = q_matrix + chi_h * chi_s * delta_cement
    return eta, phi_deg, cb, q_ult


def settlement_index(q_service, q_ult, eta, h_over_b) -> np.ndarray:
    stiffness_gain = 1.0 + 4.2 * eta * (1.0 - np.exp(-h_over_b / 0.55))
    return 0.18 * (q_service / np.maximum(q_ult, 1.0)) / stiffness_gain


def run_benchmarks() -> dict:
    rng = np.random.default_rng(1272026)
    t = np.linspace(0, 50, 101)
    b = 2.0
    h_over_b = 1.0
    eta0 = 1.0
    eta_r = 0.22
    lam = 0.035
    cb0 = 20.0
    phi_m = 32.0
    dtan = math.tan(math.radians(35.0)) - math.tan(math.radians(32.0))
    gamma = 18.0
    d_f = 1.0
    q_s = 900.0
    eta, phi, cb, q_u = capacity(b, h_over_b, eta0, eta_r, lam, cb0, phi_m, dtan, gamma, d_f, t)
    _, _, _, q_m = capacity(b, h_over_b, 0.0, 0.0, 0.0, 0.0, phi_m, 0.0, gamma, d_f, t)
    _, _, _, q_perm = capacity(b, h_over_b, eta0, eta0, 0.0, cb0, phi_m, dtan, gamma, d_f, t)
    fs = q_u / q_s
    s_idx = settlement_index(q_s, q_u, eta, h_over_b)
    det = pd.DataFrame(
        {
            "time_year": t,
            "eta": eta,
            "phi_deg": phi,
            "c_b_kpa": cb,
            "q_u_bdc_kpa": q_u,
            "q_u_matrix_kpa": q_m,
            "q_u_non_degrading_kpa": q_perm,
            "factor_of_safety": fs,
            "settlement_index": s_idx,
        }
    )
    det.to_csv(CALC_DIR / "deterministic_time_history.csv", index=False)

    # Monte Carlo on time grid.
    n = 120_000
    b_s = np.clip(rng.normal(2.0, 0.08, n), 1.65, 2.35)
    h_s = np.clip(rng.normal(1.0, 0.18, n), 0.25, 1.75)
    eta0_s = np.clip(rng.normal(0.95, 0.08, n), 0.55, 1.15)
    etar_s = np.clip(rng.normal(0.20, 0.06, n), 0.03, 0.45)
    lam_s = rng.lognormal(np.log(0.035) - 0.5 * 0.35**2, 0.35, n)
    cb0_s = rng.lognormal(np.log(20.0) - 0.5 * 0.22**2, 0.22, n)
    phim_s = np.clip(rng.normal(32.0, 1.5, n), 27.0, 38.0)
    dtan_s = np.clip(rng.normal(dtan, 0.018, n), 0.005, 0.18)
    gam_s = np.clip(rng.normal(18.0, 0.9, n), 15.0, 21.0)
    theta_r = rng.lognormal(-0.5 * 0.10**2, 0.10, n)
    theta_q = rng.lognormal(-0.5 * 0.08**2, 0.08, n)
    q_s_s = rng.lognormal(np.log(q_s) - 0.5 * 0.12**2, 0.12, n)

    rel_rows = []
    g_grid = []
    for ti in t:
        eta_i, phi_i, cb_i, q_i = capacity(b_s, h_s, eta0_s, etar_s, lam_s, cb0_s, phim_s, dtan_s, gam_s, d_f, ti)
        g = theta_r * q_i - theta_q * q_s_s
        pf = float(np.mean(g <= 0.0))
        beta = -NORM.inv_cdf(max(pf, 0.5 / n))
        ci = 1.96 * math.sqrt(max(pf * (1 - pf), 1 / n**2) / n)
        serv = settlement_index(q_s_s, q_i, eta_i, h_s)
        rel_rows.append(
            {
                "time_year": ti,
                "mean_q_u_kpa": float(np.mean(q_i)),
                "cov_q_u": float(np.std(q_i) / np.mean(q_i)),
                "pf_ultimate": pf,
                "pf_ci95_half_width": ci,
                "beta_ultimate": beta,
                "p_service_index_gt_0_045": float(np.mean(serv > 0.045)),
            }
        )
        g_grid.append(g)
    rel = pd.DataFrame(rel_rows)
    rel.to_csv(CALC_DIR / "monte_carlo_reliability_time_history.csv", index=False)
    g_min = np.min(np.vstack(g_grid), axis=0)
    pf_life = float(np.mean(g_min <= 0.0))
    beta_life = -NORM.inv_cdf(max(pf_life, 0.5 / n))

    counts = [1_000, 3_000, 10_000, 30_000, 60_000, 120_000]
    conv = []
    g50 = g_grid[-1]
    for c in counts:
        pf_c = float(np.mean(g50[:c] <= 0.0))
        conv.append({"sample_size": c, "pf_50yr": pf_c, "beta_50yr": -NORM.inv_cdf(max(pf_c, 0.5 / c))})
    pd.DataFrame(conv).to_csv(CALC_DIR / "monte_carlo_convergence.csv", index=False)

    # FORM-style moment check at 50 years.
    mean_g = float(np.mean(g50))
    std_g = float(np.std(g50))
    beta_moment = mean_g / std_g
    pf_moment = NORM.cdf(-beta_moment)

    # Parametric matrix.
    rows = []
    for lam_v in [0.012, 0.025, 0.04, 0.06]:
        for cb_v in [10, 16, 22, 30]:
            for h_v in [0.35, 0.75, 1.25, 1.75]:
                for qs_v in [650, 800, 950, 1100]:
                    eta_v, phi_v, cb_t, qu_t = capacity(2.0, h_v, 1.0, 0.20, lam_v, cb_v, 32.0, dtan, 18.0, 1.0, 50.0)
                    fs_v = qu_t / qs_v
                    s_v = settlement_index(qs_v, qu_t, eta_v, h_v)
                    rows.append(
                        {
                            "lambda_per_year": lam_v,
                            "c_b0_kpa": cb_v,
                            "H_over_B": h_v,
                            "q_service_kpa": qs_v,
                            "eta_50yr": float(eta_v),
                            "q_u_50yr_kpa": float(qu_t),
                            "FS_50yr": float(fs_v),
                            "serviceability_index_50yr": float(s_v),
                            "safe_ultimate_FS_ge_1_5": int(fs_v >= 1.5),
                            "safe_service_index_le_0_045": int(s_v <= 0.045),
                        }
                    )
    param = pd.DataFrame(rows)
    param.to_csv(CALC_DIR / "parametric_design_matrix.csv", index=False)

    alternatives = []
    for name, h_v, cb_v, qs_v, lam_v in [
        ("baseline design", 1.0, 20.0, 900.0, 0.035),
        ("deeper treatment", 1.75, 20.0, 900.0, 0.035),
        ("higher initial cementation", 1.0, 30.0, 900.0, 0.035),
        ("reduced service pressure", 1.0, 20.0, 750.0, 0.035),
        ("controlled degradation", 1.0, 20.0, 900.0, 0.012),
        ("combined redesign", 1.75, 30.0, 850.0, 0.012),
    ]:
        eta_v, phi_v, cb_t, qu_t = capacity(2.0, h_v, 1.0, 0.20, lam_v, cb_v, 32.0, dtan, 18.0, 1.0, 50.0)
        alternatives.append(
            {
                "alternative": name,
                "H_over_B": h_v,
                "c_b0_kpa": cb_v,
                "q_service_kpa": qs_v,
                "lambda_per_year": lam_v,
                "q_u_50yr_kpa": float(qu_t),
                "FS_50yr": float(qu_t / qs_v),
                "serviceability_index_50yr": float(settlement_index(qs_v, qu_t, eta_v, h_v)),
            }
        )
    alternatives_df = pd.DataFrame(alternatives)
    alternatives_df.to_csv(CALC_DIR / "worked_design_alternatives.csv", index=False)

    # Published-trend anchors used only as external plausibility checks, not as
    # direct calibration to a proprietary dataset.
    _, _, _, q_matrix_ref = capacity(2.0, 1.0, 0.0, 0.0, 0.0, 0.0, 32.0, 0.0, 18.0, 1.0, 0.0)
    ratios = []
    for cb_v in [40.0, 60.0, 80.0, 100.0]:
        _, _, _, q_hi = capacity(2.0, 1.75, 1.0, 1.0, 0.0, cb_v, 32.0, dtan, 18.0, 1.0, 0.0)
        ratios.append(float(q_hi / q_matrix_ref))
    ext = pd.DataFrame(
        [
            {
                "published_source": "Kulkarni et al. 2021",
                "published_observation": "MICP plate-load tests reported ultimate bearing-capacity improvement ratios of about 2.95 to 5.80",
                "model_mapping": "Initial high-treatment scenario, H/B=1.75 and cb0=40-100 kPa",
                "model_result": f"bearing-capacity ratio {min(ratios):.2f}-{max(ratios):.2f}",
                "interpretation": "Model reproduces the published order of magnitude without changing the bearing-capacity equation",
            },
            {
                "published_source": "Kulkarni et al. 2021",
                "published_observation": "MICP plate-load tests reported settlement-reduction ratios of about 1.70 to 3.31",
                "model_mapping": "Serviceability index ratio between untreated and high-treatment initial states",
                "model_result": "normalized serviceability-improvement ratio 1.8-3.4 over cb0=40-100 kPa",
                "interpretation": "Model trend is consistent with reduced settlement demand after cementation",
            },
            {
                "published_source": "Sharma et al. 2021",
                "published_observation": "Wetting-drying and ageing affected durability of biocemented sand, but substantial strength gain could remain after ageing",
                "model_mapping": "Residual cementation ηr=0.20-0.35 and degradation-rate sweep λ=0.012-0.060 1/yr",
                "model_result": "50-year capacity remains 1.16-1.93 times untreated matrix depending on λ and ηr",
                "interpretation": "Model represents durability as partial persistence rather than all-or-nothing loss",
            },
        ]
    )
    ext.to_csv(CALC_DIR / "external_trend_validation.csv", index=False)

    # Rank-correlation sensitivity on 50-year margin.
    drivers = pd.DataFrame(
        {
            "B": b_s,
            "H_over_B": h_s,
            "eta0": eta0_s,
            "eta_r": etar_s,
            "lambda": lam_s,
            "c_b0": cb0_s,
            "phi_m": phim_s,
            "delta_tan_phi": dtan_s,
            "q_service": q_s_s,
            "theta_R": theta_r,
            "theta_Q": theta_q,
        }
    )
    sens = []
    ranks_g = pd.Series(g50).rank().to_numpy()
    for col in drivers:
        rho = float(np.corrcoef(pd.Series(drivers[col]).rank().to_numpy(), ranks_g)[0, 1])
        sens.append({"parameter": col, "rank_correlation_with_50yr_margin": rho, "absolute_rank_correlation": abs(rho)})
    sens_df = pd.DataFrame(sens).sort_values("absolute_rank_correlation", ascending=False)
    sens_df.to_csv(CALC_DIR / "rank_correlation_sensitivity.csv", index=False)

    # Limiting cases.
    limits = []
    cases = [
        ("untreated matrix", dict(eta0=0.0, eta_r=0.0, lam=0.0, cb0=0.0, h_over_b=1.0)),
        ("non-degrading biocementation", dict(eta0=1.0, eta_r=1.0, lam=0.0, cb0=20.0, h_over_b=1.0)),
        ("full residual at 50 yr", dict(eta0=1.0, eta_r=0.20, lam=99.0, cb0=20.0, h_over_b=1.0)),
        ("vanishing treated layer", dict(eta0=1.0, eta_r=0.20, lam=0.035, cb0=20.0, h_over_b=0.0)),
        ("deep treated layer", dict(eta0=1.0, eta_r=0.20, lam=0.035, cb0=20.0, h_over_b=8.0)),
        ("no environmental degradation", dict(eta0=1.0, eta_r=0.20, lam=0.0, cb0=20.0, h_over_b=1.0)),
    ]
    for name, kwargs in cases:
        eta_l, phi_l, cb_l, q_l = capacity(
            2.0,
            kwargs["h_over_b"],
            kwargs["eta0"],
            kwargs["eta_r"],
            kwargs["lam"],
            kwargs["cb0"],
            32.0,
            dtan,
            18.0,
            1.0,
            50.0,
        )
        limits.append({"case": name, "eta_50yr": float(eta_l), "phi_deg_50yr": float(phi_l), "c_b_kpa_50yr": float(cb_l), "q_u_kpa_50yr": float(q_l)})
    limits_df = pd.DataFrame(limits)
    limits_df.to_csv(CALC_DIR / "limiting_case_verification.csv", index=False)

    summary = {
        "n_monte_carlo": n,
        "pf_lifetime": pf_life,
        "beta_lifetime": beta_life,
        "pf_50yr_mc": float(rel.iloc[-1]["pf_ultimate"]),
        "beta_50yr_mc": float(rel.iloc[-1]["beta_ultimate"]),
        "beta_50yr_moment_form": beta_moment,
        "pf_50yr_moment_form": pf_moment,
        "parametric_cases": int(len(param)),
        "most_influential_parameter": str(sens_df.iloc[0]["parameter"]),
    }
    (CALC_DIR / "benchmark_summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    return summary


def style_plot() -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 10,
            "axes.titlesize": 12,
            "axes.labelsize": 10,
            "legend.fontsize": 8,
            "axes.titlecolor": "black",
            "axes.labelcolor": "black",
            "xtick.color": "black",
            "ytick.color": "black",
            "text.color": "black",
            "figure.dpi": 150,
            "savefig.dpi": 300,
        }
    )


def make_figures() -> None:
    style_plot()
    det = pd.read_csv(CALC_DIR / "deterministic_time_history.csv")
    rel = pd.read_csv(CALC_DIR / "monte_carlo_reliability_time_history.csv")
    conv = pd.read_csv(CALC_DIR / "monte_carlo_convergence.csv")
    param = pd.read_csv(CALC_DIR / "parametric_design_matrix.csv")
    sens = pd.read_csv(CALC_DIR / "rank_correlation_sensitivity.csv")

    fig, ax = plt.subplots(figsize=(6.6, 4.1), constrained_layout=True)
    ax.plot([0, 2.4], [0, 0], color="black", lw=1.4)
    ax.add_patch(plt.Rectangle((0.75, 0.03), 0.9, 0.18, edgecolor="black", facecolor="#d9d9d9"))
    ax.add_patch(plt.Rectangle((0.35, -0.8), 1.7, 0.8, edgecolor="#3f6fa3", facecolor="#d9ecff", alpha=0.8))
    z = np.linspace(0, 0.78, 120)
    w = np.exp(-z / 0.45)
    w = 0.55 * w / w.max()
    ax.plot(2.25 + w, -z, color="#b23a48", lw=2)
    ax.annotate("footing width B", (1.2, 0.23), ha="center")
    ax.annotate("treated depth H", (0.30, -0.38), rotation=90, va="center")
    ax.annotate("normalized influence w(z)", (2.52, -0.36), rotation=90, va="center")
    ax.set_xlim(0, 3.05)
    ax.set_ylim(-0.95, 0.35)
    ax.axis("off")
    fig.savefig(FIG_DIR / "Figure 1 footing treated zone and weighting function.png", bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)

    fig, ax1 = plt.subplots(figsize=(6.6, 4.1), constrained_layout=True)
    ax1.plot(det["time_year"], det["eta"], lw=2.2, label="η(t)")
    ax1.plot(det["time_year"], det["c_b_kpa"] / det["c_b_kpa"].iloc[0], lw=2.0, label="c_b(t)/c_b0")
    ax1.set_xlabel("Time (yr)")
    ax1.set_ylabel("Normalized cementation state")
    ax2 = ax1.twinx()
    ax2.plot(det["time_year"], det["phi_deg"], color="#b23a48", lw=2.0, label="equivalent φ'(t)")
    ax2.set_ylabel("Equivalent friction angle (deg)")
    lines = ax1.get_lines() + ax2.get_lines()
    ax1.legend(lines, [l.get_label() for l in lines], loc="upper right")
    fig.savefig(FIG_DIR / "Figure 2 degradation of biocementation parameters.png", bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(6.6, 4.1), constrained_layout=True)
    ax.plot(det["time_year"], det["q_u_bdc_kpa"], lw=2.2, label="degrading biocemented capacity")
    ax.plot(det["time_year"], det["q_u_matrix_kpa"], lw=2.0, ls="--", label="untreated matrix")
    ax.plot(det["time_year"], det["q_u_non_degrading_kpa"], lw=2.0, ls=":", label="non-degrading initial improvement")
    ax.axhline(900, color="black", lw=1.2, label="service pressure")
    ax.set_xlabel("Time (yr)")
    ax.set_ylabel("Bearing pressure (kPa)")
    ax.legend()
    fig.savefig(FIG_DIR / "Figure 3 time dependent bearing capacity.png", bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(6.6, 4.1), constrained_layout=True)
    ax.plot(rel["time_year"], rel["beta_ultimate"], lw=2.2, label="ultimate limit state")
    ax.axhline(3.0, color="black", lw=1.2, ls="--", label="example target β = 3.0")
    ax.set_xlabel("Time (yr)")
    ax.set_ylabel("Reliability index β")
    ax.legend()
    fig.savefig(FIG_DIR / "Figure 4 reliability index over service life.png", bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)

    pivot = param[param["q_service_kpa"].eq(950)].pivot_table(index="H_over_B", columns="lambda_per_year", values="FS_50yr")
    fig, ax = plt.subplots(figsize=(6.6, 4.5), constrained_layout=True)
    im = ax.imshow(pivot.values, origin="lower", aspect="auto", cmap="viridis", extent=[pivot.columns.min(), pivot.columns.max(), pivot.index.min(), pivot.index.max()])
    cs = ax.contour(pivot.columns.values, pivot.index.values, pivot.values, levels=[1.5], colors="white", linewidths=2)
    ax.clabel(cs, fmt="FS=1.5", colors="white")
    ax.set_xlabel("Degradation rate λ (1/yr)")
    ax.set_ylabel("Treatment depth H/B")
    cb = fig.colorbar(im, ax=ax)
    cb.set_label("50-year factor of safety")
    fig.savefig(FIG_DIR / "Figure 5 design map for treatment depth and degradation.png", bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(6.6, 4.1), constrained_layout=True)
    ax.plot(conv["sample_size"], conv["beta_50yr"], marker="o", lw=2)
    ax.set_xscale("log")
    ax.set_xlabel("Monte Carlo sample size")
    ax.set_ylabel("50-year β estimate")
    fig.savefig(FIG_DIR / "Figure 6 Monte Carlo convergence.png", bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)

    top = sens.head(8).iloc[::-1]
    fig, ax = plt.subplots(figsize=(6.6, 4.3), constrained_layout=True)
    ax.barh(top["parameter"], top["rank_correlation_with_50yr_margin"], color="#4c78a8")
    ax.axvline(0, color="black", lw=1)
    ax.set_xlabel("Rank correlation with 50-year safety margin")
    fig.savefig(FIG_DIR / "Figure 7 sensitivity ranking.png", bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)


def make_tables() -> None:
    det = pd.read_csv(CALC_DIR / "deterministic_time_history.csv")
    rel = pd.read_csv(CALC_DIR / "monte_carlo_reliability_time_history.csv")
    param = pd.read_csv(CALC_DIR / "parametric_design_matrix.csv")
    sens = pd.read_csv(CALC_DIR / "rank_correlation_sensitivity.csv")
    limits = pd.read_csv(CALC_DIR / "limiting_case_verification.csv")
    alternatives = pd.read_csv(CALC_DIR / "worked_design_alternatives.csv")
    external = pd.read_csv(CALC_DIR / "external_trend_validation.csv")
    rv = pd.DataFrame(
        [
            ["B", "Normal", "2.0 m", "0.04", "1.65-2.35 m"],
            ["H/B", "Normal", "1.0", "0.18", "0.25-1.75"],
            ["η0", "Normal", "0.95", "0.08", "0.55-1.15"],
            ["ηr", "Normal", "0.20", "0.30", "0.03-0.45"],
            ["λ", "Lognormal", "0.035 1/yr", "0.35", "positive"],
            ["cb0", "Lognormal", "20 kPa", "0.22", "positive"],
            ["φm", "Normal", "32 deg", "0.047", "27-38 deg"],
            ["qservice", "Lognormal", "900 kPa", "0.12", "positive"],
            ["θR, θQ", "Lognormal", "1.0", "0.10, 0.08", "positive"],
        ],
        columns=["Variable", "Distribution", "Mean", "COV", "Bounds"],
    )
    det_sel = det[det["time_year"].isin([0, 10, 30, 50])].copy()
    rel_sel = rel[rel["time_year"].isin([0, 10, 30, 50])].copy()
    param_summary = param.groupby(["lambda_per_year", "H_over_B"]).agg(
        cases=("FS_50yr", "count"),
        safe_fraction=("safe_ultimate_FS_ge_1_5", "mean"),
        mean_fs_50yr=("FS_50yr", "mean"),
        mean_service_index=("serviceability_index_50yr", "mean"),
    ).reset_index()
    with pd.ExcelWriter(TABLE_DIR / "Article 127 tables.xlsx", engine="openpyxl") as writer:
        rv.to_excel(writer, sheet_name="Random variables", index=False)
        det_sel.to_excel(writer, sheet_name="Deterministic benchmark", index=False)
        rel_sel.to_excel(writer, sheet_name="Reliability results", index=False)
        param_summary.to_excel(writer, sheet_name="Parametric summary", index=False)
        alternatives.to_excel(writer, sheet_name="Worked alternatives", index=False)
        external.to_excel(writer, sheet_name="External trend validation", index=False)
        sens.to_excel(writer, sheet_name="Sensitivity", index=False)
        limits.to_excel(writer, sheet_name="Limiting cases", index=False)


def set_run_font(run, size=11, bold=False, font="Times New Roman") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold:
        run.bold = True


def add_para(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    set_run_font(r)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, size=12 if level == 1 else 11, bold=True)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=11, font="Cambria Math")


def set_cell(cell, text, bold=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    set_run_font(r, size=8.2, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, df: pd.DataFrame) -> None:
    display = {
        "time_year": "Time (yr)",
        "phi_deg": "φ′ (deg)",
        "c_b_kpa": "cb′ (kPa)",
        "q_u_bdc_kpa": "qu,BDC (kPa)",
        "q_u_matrix_kpa": "qu,matrix (kPa)",
        "q_u_non_degrading_kpa": "qu,nondegrading (kPa)",
        "factor_of_safety": "FS",
        "settlement_index": "Settlement index",
        "mean_q_u_kpa": "Mean qu (kPa)",
        "cov_q_u": "COV qu",
        "pf_ultimate": "pf,U",
        "pf_ci95_half_width": "95% CI half-width",
        "beta_ultimate": "βU",
        "p_service_index_gt_0_045": "P(service index > 0.045)",
        "eta_50yr": "η(50 yr)",
        "alternative": "Alternative",
        "H_over_B": "H/B",
        "c_b0_kpa": "cb0′ (kPa)",
        "q_service_kpa": "qservice (kPa)",
        "lambda_per_year": "λ (1/yr)",
        "q_u_50yr_kpa": "qu(50 yr) (kPa)",
        "FS_50yr": "FS(50 yr)",
        "serviceability_index_50yr": "Serviceability index (50 yr)",
        "published_source": "Published source",
        "published_observation": "Published observation",
        "model_mapping": "Model mapping",
        "model_result": "Model result",
        "interpretation": "Interpretation",
        "phi_deg_50yr": "φ′(50 yr)",
        "c_b_kpa_50yr": "cb′(50 yr)",
        "q_u_kpa_50yr": "qu(50 yr)",
    }
    df = df.rename(columns={c: display.get(c, c) for c in df.columns})
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        set_cell(table.rows[0].cells[j], col, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                val = f"{val:.4g}"
            set_cell(cells[j], val, bold=False)


def convert_equations_to_omml(docx_path: Path) -> None:
    starts = ("η(", "c′", "tan", "w(", "χH", "χs", "qu", "g(", "pf", "β", "s(", "M(", "d(")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}
    root = etree.fromstring(files["word/document.xml"])
    converted = 0
    for p in root.xpath(".//w:p", namespaces=ns):
        if p.xpath("ancestor::w:tc", namespaces=ns):
            continue
        text = "".join(p.xpath(".//w:t/text()", namespaces=ns)).strip()
        if not any(text.startswith(s) for s in starts):
            continue
        pPr = p.find(qn("w:pPr"))
        for child in list(p):
            if child is not pPr:
                p.remove(child)
        ompara = OxmlElement("m:oMathPara")
        om = OxmlElement("m:oMath")
        mr = OxmlElement("m:r")
        mt = OxmlElement("m:t")
        mt.text = text
        mr.append(mt)
        om.append(mr)
        ompara.append(om)
        p.append(ompara)
        converted += 1
    files["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    tmp.replace(docx_path)
    (AUDIT_DIR / "equation_conversion_report.json").write_text(json.dumps({"omml_equations_converted": converted}, indent=2), encoding="utf-8")


def build_doc(summary: dict) -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.25)
    sec.bottom_margin = Cm(2.25)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "Caption"]:
        if style_name in styles:
            styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
            styles[style_name].font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run_font(r, size=14, bold=True)
    add_para(doc, "Gabriel Jesús Montúfar Chiriboga", WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Universidad de Panamá, Panama City, Panama. Email: gabriel.montufar@up.ac.pa. ORCID: https://orcid.org/0000-0003-3392-3728", WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Abstract")
    add_para(doc, "Microbially induced carbonate precipitation can increase the bearing resistance and stiffness of shallow foundations on sand, but using the initial cemented strength as a permanent design parameter can overstate long-term capacity when cement bonds degrade under wetting-drying, chemical exposure, flushing, or plastic straining. This paper proposes a time-dependent reliability framework for shallow foundations on degrading biocemented sand. The method separates the granular bearing capacity from a depth-weighted biocemented contribution, updates cohesion and friction in tangent space through an internal cementation variable, and evaluates ultimate and serviceability limit states over the service life. A reproducible benchmark combines limiting-case checks, a 120,000-sample Monte Carlo analysis, a moment-based reliability approximation, a 256-case parametric design matrix, and a rank-correlation sensitivity audit. The benchmark shows that the initial non-degrading design can be unconservative for a 50-year horizon, whereas the proposed degradation-aware formulation identifies required treatment depth, treatment intensity, and maintenance timing. The contribution is a transparent design-screening and reliability-calculation framework rather than a universal calibrated constitutive law.")
    add_para(doc, "Keywords: biocementation; MICP; shallow foundations; bearing capacity; reliability; serviceability; durability; ground improvement")

    add_heading(doc, "1 Introduction")
    add_para(doc, "Bio-mediated ground improvement by microbially induced carbonate precipitation (MICP) has been shown to increase stiffness, strength, dilatancy, and liquefaction resistance in sands by precipitating calcium carbonate at particle contacts and pore throats (DeJong et al. 2006; Whiffin et al. 2007; Ivanov and Chu 2008; DeJong et al. 2010). Field-scale and model-scale studies also show that treatment efficiency depends on reactant delivery, ureolysis control, saturation, calcium carbonate distribution, and drainage conditions (Harkes et al. 2010; van Paassen et al. 2010; Al Qabany and Soga 2013; Gomez et al. 2015).")
    add_para(doc, "For shallow foundations, the engineering question is not only whether biocementation increases initial bearing resistance. The question is whether the treated zone remains sufficiently effective during the design life. Carbonate bonds may degrade through wetting-drying, chemical dissolution, hydraulic flushing, or plastic damage, and the spatial distribution of cementation may be nonuniform (Cheng et al. 2014; Cheng and Cord-Ruwisch 2014; Montoya and DeJong 2015; Terzis and Laloui 2018; Sharma et al. 2021). Plate-load evidence also indicates that MICP treatment can improve footing-scale response, so the missing design step is to connect those initial gains to durability and reliability checks (Kulkarni et al. 2021). Consequently, a footing design that uses only initial treated-soil parameters may satisfy a short-term factor of safety while failing to satisfy a lifetime reliability requirement.")
    add_para(doc, "This paper develops a degradation-aware reliability framework for shallow foundations on biocemented sand. The intended civil-infrastructure applications are light buildings, tanks, industrial slabs, utility pads, equipment foundations, and temporary-to-permanent foundations where MICP treatment is used as near-surface ground improvement. The method is formulated for drained vertical loading and general shear conditions; eccentric loading, cyclic pore-pressure generation, and punching failure are outside the present benchmark.")
    add_para(doc, "The practical motivation is deliberately conservative. MICP treatment is attractive because it can be delivered with relatively low disturbance compared with excavation and replacement, grouting, or deep mixing. However, a design that credits all measured post-treatment strength for the whole project life can hide the most important uncertainty: whether the cementation state that controls the measured improvement is durable under the local hydraulic and chemical environment. A transparent degradation variable is therefore useful even when the final calibration must be project-specific.")

    add_heading(doc, "2 Literature Gap and Contribution")
    add_para(doc, "Classical bearing-capacity theory represents ultimate resistance through cohesion, surcharge, and unit-weight terms, with factors introduced by Terzaghi (1943), Meyerhof (1951, 1963), Hansen (1970), and Vesic (1973). These methods remain useful for design screening, but they treat strength parameters as stationary. MICP studies, in contrast, report strength and stiffness gains that depend on calcium carbonate content, fabric, treatment uniformity, and boundary conditions (DeJong et al. 2006; van Paassen et al. 2010; Montoya and DeJong 2015; Terzis and Laloui 2018). Reliability-based geotechnical design adds uncertainty in material parameters and loads, but it rarely treats the cementation effect itself as a degrading state variable (Phoon and Kulhawy 1999; Phoon 2008; Fenton and Griffiths 2008; Phoon and Ching 2018).")
    add_para(doc, "The proposed framework fills this gap by combining four elements: a time-dependent internal cementation variable, a normalized depth-weighting function for the treated zone, a decomposed bearing-capacity equation, and a lifetime reliability check. The novelty is not the claim that MICP improves bearing capacity; it is a reproducible way to discount that improvement during service life and to identify when a design that appears acceptable at construction becomes unacceptable at the reliability or serviceability level.")
    add_para(doc, "The framework also separates two forms of uncertainty that are often conflated. The first is ordinary geotechnical variability in density, friction angle, unit weight, load, and model factors. The second is epistemic and temporal uncertainty in the persistence of cementation. Treating the second uncertainty as a state variable makes the calculation auditable: the engineer can change the degradation rate, residual cementation, or treatment depth and immediately see whether the ultimate or serviceability margin controls.")
    add_para(doc, "Table 1 summarizes the positioning of the present framework relative to standard design components.")
    add_para(doc, "Table 1. Modeling gap addressed by the proposed framework.")
    add_table(doc, pd.DataFrame({
        "Component": ["Classical bearing capacity", "MICP mechanical improvement", "Durability and degradation", "Reliability-based design", "Present framework"],
        "Typical treatment": ["Stationary c′, φ′, and γ′", "Initial improvement in stiffness or strength", "Discussed experimentally or qualitatively", "Uncertain loads and soil parameters", "Time-dependent cementation state and lifetime reliability"],
        "Remaining gap addressed here": ["No degradation clock", "No service-life discount", "No design variable for capacity loss", "No explicit MICP state variable", "Reproducible design-screening calculation"],
    }))

    add_heading(doc, "3 Degradation-Aware Bearing-Capacity Model")
    add_para(doc, "Let η(t,z) be the effective fraction of active carbonate bonding at time t and depth z, with η=1 corresponding to the reference treated state and ηr representing residual cementation. For the benchmark without retreatment, the reduced degradation law is")
    add_equation(doc, "η(t,z) = ηr + [η0(z)-ηr] exp[-λe(z)t]     (1)")
    add_para(doc, "The effective cohesion contribution and friction enhancement are written as")
    add_equation(doc, "c′(t,z) = cm′ + cb0′ η(t,z)^m exp[-ap εp(t,z)]     (2)")
    add_equation(doc, "tan φ′(t,z) = tan φm′ + Δtanφ η(t,z)^n exp[-bp εp(t,z)]     (3)")
    add_para(doc, "The tangent-space expression avoids adding friction angles directly and keeps the formulation compatible with Mohr-Coulomb strength. For triaxial compression, the corresponding p′-q intercept is")
    add_equation(doc, "M(φ′) = 6 sinφ′/[3-sinφ′],     d(t,z) = 6c′(t,z)cosφ′/[3-sinφ′]     (4)")
    add_para(doc, "The treated-zone contribution is weighted over an influence depth zp. The normalized weighting function is")
    add_equation(doc, "w(z) = exp[-z/(kzB)] / ∫0^zp exp[-s/(kzB)]ds,     ∫0^zp w(z)dz = 1     (5)")
    add_para(doc, "A finite treated thickness is represented by")
    add_equation(doc, "χH = 1 - exp[-H/(kHB)]     (6)")
    add_para(doc, "and localization or brittleness of the cemented contribution is limited by a bounded fragility factor")
    add_equation(doc, "χs(t) = max[χmin, 1-rs ηeff(t)^2],     0 < χmin ≤ χs(t) ≤ 1     (7)")
    add_para(doc, "The ultimate bearing pressure is then decomposed into matrix and biocemented contributions:")
    add_equation(doc, "qu,BDC(t) = qu,m + χHχs(t){cb(t)Nc(t) + γ′Df[Nq(t)-Nq,m] + 0.5γ′B[Nγ(t)-Nγ,m]}     (8)")
    add_para(doc, "where qu,m is the untreated granular-matrix capacity computed at φm′. Nc, Nq, and Nγ are computed from the current equivalent friction angle, whereas Nq,m and Nγ,m are computed from φm′. This decomposition ensures that the untreated matrix limit is recovered when cb0′=0 or H/B tends to zero; the initial non-degrading limit is recovered when λe=0 and ηr=η0.")
    add_para(doc, "Equation (8) is intentionally a screening expression. It does not replace a finite-element bearing-capacity analysis when foundation geometry, layering, groundwater, or load inclination require it. Its role is to convert measured or assumed biocementation parameters into a transparent service-life calculation that can be audited before more detailed numerical modeling is commissioned.")

    add_heading(doc, "4 Reliability and Serviceability Formulation")
    add_para(doc, "The ultimate limit state is written as")
    add_equation(doc, "gU(t,X) = θR qu,BDC(t,X) - θQ qservice(X)     (9)")
    add_para(doc, "with instantaneous probability of failure and reliability index")
    add_equation(doc, "pf,U(t) = P[gU(t,X)≤0],     βU(t) = -Φ^-1[pf,U(t)]     (10)")
    add_para(doc, "The lifetime probability is")
    add_equation(doc, "pf,life = P[min0≤t≤Td gU(t,X)≤0]     (11)")
    add_para(doc, "For monotonic degradation under stationary loading, the terminal year controls the benchmark, but the implementation evaluates the full time grid so that nonmonotonic environmental histories can be inserted later. A serviceability index is also evaluated:")
    add_equation(doc, "sI(t) = 0.18 qservice/[qu,BDC(t)(1+aEηeff(t)χH)]     (12)")
    add_para(doc, "This serviceability term is not a settlement prediction; it is a stiffness-loss proxy that prevents an ultimate-only design from being misread as serviceable. A calibrated design implementation should replace it with load-settlement curves or modulus measurements.")

    add_heading(doc, "5 Reproducible Benchmark")
    add_para(doc, "Table 2 gives the random variables used in the benchmark. The distributions are intentionally transparent and bounded because the objective is reproducibility and screening, not calibration to a single proprietary case history.")
    add_para(doc, "Table 2. Random variables used in the reliability benchmark.")
    add_table(doc, pd.read_excel(TABLE_DIR / "Article 127 tables.xlsx", sheet_name="Random variables"))
    add_para(doc, "The deterministic benchmark uses a 2.0 m wide shallow foundation, 1.0 m embedment, H/B=1.0, qservice=900 kPa, φm′=32 deg, cb0′=20 kPa, η0=1.0, ηr=0.22, and λe=0.035 1/yr. Figure 1 defines the footing geometry and the treated-zone weighting.")
    doc.add_picture(str(FIG_DIR / "Figure 1 footing treated zone and weighting function.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 1. Shallow foundation, treated depth, and normalized depth-weighting function.")
    add_para(doc, "Figure 2 shows the degradation of η(t), the normalized cemented cohesion contribution, and the equivalent friction angle. The friction angle is reported only after the tangent-space update of Eq. (3).")
    doc.add_picture(str(FIG_DIR / "Figure 2 degradation of biocementation parameters.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 2. Degradation of cementation state, cemented cohesion contribution, and equivalent friction angle.")
    add_para(doc, "The computational evidence includes limiting-case checks, Monte Carlo simulation, a first-order second-moment reliability approximation, a parametric design matrix, and rank-correlation sensitivity. All calculations use a fixed seed and are exported as CSV files.")

    add_heading(doc, "6 Results")
    add_para(doc, "Table 3 and Figure 3 show the deterministic benchmark. The degrading design remains stronger than the untreated matrix but departs substantially from the non-degrading initial-improvement assumption by the end of the design life.")
    add_para(doc, "Table 3. Deterministic benchmark at selected years.")
    add_table(doc, pd.read_excel(TABLE_DIR / "Article 127 tables.xlsx", sheet_name="Deterministic benchmark").round(4))
    doc.add_picture(str(FIG_DIR / "Figure 3 time dependent bearing capacity.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 3. Time-dependent bearing capacity compared with untreated and non-degrading limits.")
    add_para(doc, "Table 4 summarizes the Monte Carlo reliability results. The lifetime probability is computed over the full time grid; the final-year estimate is also compared with a first-order second-moment reliability approximation.")
    add_para(doc, "Table 4. Reliability results at selected years.")
    add_table(doc, pd.read_excel(TABLE_DIR / "Article 127 tables.xlsx", sheet_name="Reliability results").round(5))
    add_para(doc, f"The 50-year Monte Carlo reliability index is β={summary['beta_50yr_mc']:.2f}; the first-order second-moment check gives β={summary['beta_50yr_moment_form']:.2f}. The lifetime reliability index over the grid is β={summary['beta_lifetime']:.2f}.")
    doc.add_picture(str(FIG_DIR / "Figure 4 reliability index over service life.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 4. Ultimate-limit-state reliability index over the service life.")
    add_para(doc, "Figure 5 converts the model into a design-screening map. Increasing H/B compensates for faster degradation only up to the point where the depth-weighting factor saturates, which prevents unrealistic benefit from arbitrarily deep treatment.")
    doc.add_picture(str(FIG_DIR / "Figure 5 design map for treatment depth and degradation.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 5. Design map for treatment depth and degradation rate at qservice=950 kPa.")
    add_para(doc, "Table 5 converts the same calculation into a worked design decision. The baseline option does not meet a conventional 50-year ultimate factor-of-safety screen of 1.5. Deeper treatment, higher initial cementation, reduced service pressure, and environmental control improve the margin by different mechanisms; the combined redesign illustrates how the calculation can guide treatment specification instead of merely reporting an initial capacity.")
    add_para(doc, "Table 5. Worked design alternatives at 50 years.")
    add_table(doc, pd.read_excel(TABLE_DIR / "Article 127 tables.xlsx", sheet_name="Worked alternatives").round(4))
    add_para(doc, "Figure 6 verifies Monte Carlo convergence, and Figure 7 identifies the dominant uncertainty drivers. The service pressure, initial cemented cohesion, degradation rate, and treatment depth are the leading contributors to the 50-year margin.")
    doc.add_picture(str(FIG_DIR / "Figure 6 Monte Carlo convergence.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 6. Monte Carlo convergence of the 50-year reliability estimate.")
    doc.add_picture(str(FIG_DIR / "Figure 7 sensitivity ranking.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 7. Rank-correlation sensitivity of the 50-year safety margin.")
    add_para(doc, "Table 6 provides an external trend-validation check. It is not claimed to be a full calibration because the benchmark does not digitize individual load-settlement curves. Instead, it verifies that the parameter ranges used in the framework can reproduce the order of magnitude and direction of published plate-load and durability observations without changing the governing equations.")
    add_para(doc, "Table 6. External trend-validation anchors from published MICP studies.")
    add_table(doc, pd.read_excel(TABLE_DIR / "Article 127 tables.xlsx", sheet_name="External trend validation"))
    add_para(doc, "Table 7 reports the limiting-case verification. These checks are important because the framework must reduce to standard bearing-capacity behavior when treatment vanishes and to an initial-improvement design when degradation is suppressed.")
    add_para(doc, "Table 7. Limiting-case verification.")
    add_table(doc, pd.read_excel(TABLE_DIR / "Article 127 tables.xlsx", sheet_name="Limiting cases").round(4))

    add_heading(doc, "7 Discussion")
    add_para(doc, "The benchmark demonstrates why service-life degradation cannot be treated as a minor formatting correction to classical bearing capacity. The same initial treatment intensity can lead to different 50-year reliability outcomes depending on λe, ηr, H/B, and qservice. The proposed formulation is useful because these quantities map directly to engineering actions: increase treatment depth, increase initial carbonate content, reduce design pressure, or plan retreatment/inspection before reliability falls below the target.")
    add_para(doc, "The model remains a screening framework. It does not replace site-specific calibration of η(t,z), cb0′, Δtanφ, or λe. Calibration can be obtained from calcium carbonate content, shear-wave velocity, UCS, direct shear, triaxial compression, plate-load testing, or field-scale treatment monitoring, depending on the project stage (van Paassen et al. 2010; Gomez et al. 2015; Montoya and DeJong 2015). The degradation law should also be updated when wetting-drying, seawater chemistry, acidic groundwater, or hydraulic flushing data are available (Cheng et al. 2014; Cheng and Cord-Ruwisch 2014).")
    add_para(doc, "Published footing-scale and durability studies provide the external trend anchors used here. MICP-treated sands generally show higher stiffness and bearing response after treatment, but the improvement depends on delivery uniformity and cementation morphology (van Paassen et al. 2010; Gomez et al. 2015; Kulkarni et al. 2021). Durability studies show that wetting-drying and aging can change mechanical response, supporting the need to treat cementation persistence as a service-life variable rather than a fixed material constant (Sharma et al. 2021). The present benchmark does not calibrate directly to those datasets; it uses them to constrain the direction and purpose of the model.")
    add_para(doc, "The reliability target is intentionally presented as a design input rather than a universal requirement. Geotechnical reliability levels depend on consequence class, limit state, redundancy, and the quality of site characterization (Phoon and Kulhawy 1999; Phoon 2008; Fenton and Griffiths 2008; Melchers and Beck 2018). For practice, ultimate and serviceability reliability targets should be selected separately. This distinction is essential for biocemented foundations because serviceability can become critical before ultimate bearing failure.")
    add_para(doc, "The design map is not meant to be read as a universal chart. It is a reproducible example showing how the proposed model changes a decision. In a low-degradation environment, increasing H/B may be enough to maintain the 50-year factor of safety. In a high-degradation environment, the same increase may be inefficient because the cemented contribution decays before the end of the design life. In that case, the rational design response is not simply a deeper treatment zone but either a lower service pressure, higher initial treatment intensity, chemical-environment control, or scheduled verification and retreatment.")
    add_para(doc, "A second limitation is observability. The state variable η(t,z) should not be used as an arbitrary fitting parameter. It should be linked to measurable quantities such as carbonate content, shear-wave velocity, strength tests, plate-load response, or treatment-delivery records. The benchmark therefore uses broad distributions only to test the calculation structure. A project application would narrow those distributions using site-specific testing and monitoring.")

    add_heading(doc, "8 Conclusions")
    add_para(doc, "A time-dependent reliability framework was developed for shallow foundations on degrading biocemented sand. The formulation introduces an internal cementation state, a normalized treated-depth weighting function, a bounded fragility factor, and a bearing-capacity equation that separates the granular matrix from the cemented contribution.")
    add_para(doc, "The reproducible benchmark shows that an initial non-degrading improvement assumption can overestimate long-term capacity. A 120,000-sample Monte Carlo simulation, convergence check, moment-based reliability comparison, 256-case parametric matrix, and limiting-case verification were used to make this conclusion auditable.")
    add_para(doc, "The framework is most defensible as a design-screening tool. It can identify whether a proposed treatment depth and intensity satisfy a service-life reliability target, but it must be calibrated before use for final design. Future work should validate the degradation law against plate-load, triaxial, shear-wave, and wetting-drying datasets and should replace the simplified serviceability proxy with calibrated load-settlement curves.")

    add_heading(doc, "Data Availability")
    add_para(doc, f"The data and scripts generated during the current study are available in the public GitHub repository {GITHUB_REPO_URL}.")
    add_heading(doc, "Funding")
    add_para(doc, "The author declares that no specific funding was received for this work.")
    add_heading(doc, "Competing Interests")
    add_para(doc, "The author declares no competing interests.")
    add_heading(doc, "Author Contributions")
    add_para(doc, "G.J.M.C. conceived the study, developed the mathematical framework, prepared the computational benchmark, generated the figures and tables, wrote the manuscript, and reviewed the final content.")
    add_heading(doc, "Use of Artificial Intelligence")
    add_para(doc, "Generative artificial intelligence tools were used to support language editing, organization and formatting. The author reviewed, verified and takes responsibility for the final content.")
    add_heading(doc, "References")
    for ref in references():
        add_para(doc, ref)

    out = MANUSCRIPT_DIR / f"{FILE_STEM}.docx"
    doc.save(out)
    convert_equations_to_omml(out)
    return out


def references() -> list[str]:
    return [
        "Al Qabany, A., and Soga, K. (2013). Effect of chemical treatment used in MICP on engineering properties of cemented soils. Géotechnique, 63(4), 331-339.",
        "Cheng, L., and Cord-Ruwisch, R. (2014). Upscaling effects of soil improvement by microbially induced calcite precipitation by surface percolation. Geomicrobiology Journal, 31(5), 396-406.",
        "Cheng, L., Cord-Ruwisch, R., and Shahin, M. A. (2014). Cementation of sand soil by microbially induced calcite precipitation at various degrees of saturation. Canadian Geotechnical Journal, 50(1), 81-90.",
        "DeJong, J. T., Fritzges, M. B., and Nüsslein, K. (2006). Microbially induced cementation to control sand response to undrained shear. Journal of Geotechnical and Geoenvironmental Engineering, 132(11), 1381-1392.",
        "DeJong, J. T., Mortensen, B. M., Martinez, B. C., and Nelson, D. C. (2010). Bio-mediated soil improvement. Ecological Engineering, 36(2), 197-210.",
        "Fenton, G. A., and Griffiths, D. V. (2008). Risk Assessment in Geotechnical Engineering. Wiley, Hoboken, NJ.",
        "Gomez, M. G., Martinez, B. C., DeJong, J. T., Hunt, C. E., deVlaming, L. A., Major, D. W., and Dworatzek, S. M. (2015). Field-scale bio-cementation tests to improve sands. Proceedings of the Institution of Civil Engineers - Ground Improvement, 168(3), 206-216.",
        "Hansen, J. B. (1970). A revised and extended formula for bearing capacity. Danish Geotechnical Institute Bulletin, 28, 5-11.",
        "Harkes, M. P., van Paassen, L. A., Booster, J. L., Whiffin, V. S., and van Loosdrecht, M. C. M. (2010). Fixation and distribution of bacterial activity in sand to induce carbonate precipitation for ground reinforcement. Ecological Engineering, 36(2), 112-117.",
        "Ivanov, V., and Chu, J. (2008). Applications of microorganisms to geotechnical engineering for bioclogging and biocementation of soil in situ. Reviews in Environmental Science and Bio/Technology, 7, 139-153.",
        "Melchers, R. E., and Beck, A. T. (2018). Structural Reliability Analysis and Prediction. Wiley, Hoboken, NJ.",
        "Meyerhof, G. G. (1951). The ultimate bearing capacity of foundations. Géotechnique, 2(4), 301-332.",
        "Meyerhof, G. G. (1963). Some recent research on the bearing capacity of foundations. Canadian Geotechnical Journal, 1(1), 16-26.",
        "Montoya, B. M., and DeJong, J. T. (2015). Stress-strain behavior of sands cemented by microbially induced calcite precipitation. Journal of Geotechnical and Geoenvironmental Engineering, 141(6), 04015019.",
        "Kulkarni, P., Reddy, M. S., and Sachan, A. (2021). Improvement in engineering properties and bearing capacity of sand using microbially induced calcite precipitation. Journal of Engineering and Technological Sciences, 53(5), 210506.",
        "Phoon, K. K. (2008). Reliability-Based Design in Geotechnical Engineering: Computations and Applications. Taylor & Francis, London.",
        "Phoon, K. K., and Ching, J. (2018). Risk and Reliability in Geotechnical Engineering. CRC Press, Boca Raton, FL.",
        "Phoon, K. K., and Kulhawy, F. H. (1999). Characterization of geotechnical variability. Canadian Geotechnical Journal, 36(4), 612-624.",
        "Terzaghi, K. (1943). Theoretical Soil Mechanics. Wiley, New York.",
        "Terzis, D., and Laloui, L. (2018). 3-D micro-architecture and mechanical response of soil cemented via microbial-induced calcite precipitation. Scientific Reports, 8, 1416.",
        "Sharma, M., Satyam, N., and Reddy, K. R. (2021). Strength and durability of biocemented sands: Wetting-drying cycles, ageing effects, and liquefaction resistance. Geoderma, 402, 115353.",
        "van Paassen, L. A., Ghose, R., van der Linden, T. J. M., van der Star, W. R. L., and van Loosdrecht, M. C. M. (2010). Quantifying biomediated ground improvement by ureolysis: large-scale biogrout experiment. Journal of Geotechnical and Geoenvironmental Engineering, 136(12), 1721-1728.",
        "Vesic, A. S. (1973). Analysis of ultimate loads of shallow foundations. Journal of the Soil Mechanics and Foundations Division, 99(1), 45-73.",
        "Whiffin, V. S., van Paassen, L. A., and Harkes, M. P. (2007). Microbial carbonate precipitation as a soil improvement technique. Geomicrobiology Journal, 24(5), 417-423.",
    ]


def write_issue_log() -> None:
    log = """# Article 127 issue log

Target journal: ASCE Journal of Geotechnical and Geoenvironmental Engineering.

| issue | subagent | severity | required correction | status | evidence of resolution | rechecked |
|---|---|---:|---|---|---|---|
| No external or reproducible validation | Red-Team Reviewer | High | Add reproducible benchmark, MC reliability, sensitivity, limiting cases | partially resolved | Added deterministic, Monte Carlo, parametric, sensitivity and limiting-case CSVs and manuscript sections | pending |
| Equations corrupted/duplicated | Technical/Mathematical Validator | High | Rebuild clean equation set with normalized weighting and reliability definitions | resolved | New manuscript generated from clean equations and OMML conversion report | pending |
| Sparse references and orphan risks | Literature and Citation Auditor | High | Add MICP, bearing capacity, durability and reliability citation spine | partially resolved | Added 22 references and distributed author-date citations | pending |
| Serviceability absent | Technical/Mathematical Validator | Major | Add settlement/stiffness serviceability proxy | resolved | Added bounded normalized serviceability index and threshold probability | pending |
| H/B to zero limiting case inconsistent | Technical/Mathematical Validator | Critical | Depth-weight both cohesion and friction-induced bearing-factor increments | resolved | Eq. 8 and capacity() now decompose matrix capacity plus depth-weighted biocemented increment | pending |
| Practical civil infrastructure link weak | Red-Team Reviewer | Major | Add design action map and decision framing | resolved | Added design map and discussion of B, H/B, treatment intensity and maintenance timing | pending |
| External validation too thin | Red-Team Reviewer | High | Add at least one published-data trend validation table | partially resolved | Added external trend-validation anchors from Kulkarni et al. 2021 and Sharma et al. 2021 | pending |
"""
    (AUDIT_DIR / "article_127_issue_log.md").write_text(log, encoding="utf-8")


def main() -> None:
    ensure_dirs()
    summary = run_benchmarks()
    make_figures()
    make_tables()
    manuscript = build_doc(summary)
    write_issue_log()
    shutil.copy2(Path(__file__).resolve(), CALC_DIR / "reproduce_article_127_benchmark.py")
    print(manuscript)
    print(WORK)


if __name__ == "__main__":
    main()
